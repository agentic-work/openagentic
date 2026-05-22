#!/usr/bin/env python3
"""
handoff-package-bundle.py — produce the [PLATFORM] customer deliverable.

Reads every markdown file under docs/handoff-package/ and emits:
  1. A combined PDF per individual document  (via weasyprint)
  2. One unified DOCX containing every section (via python-docx)
  3. An index.html for human review of the bundle

Customer deliverable per #676. No network I/O, no platform brand leaks
(everything passes through `[PLATFORM]` substitution as authored).

Usage:
    python3 scripts/handoff-package-bundle.py
"""
from __future__ import annotations

import re
import subprocess
import sys
from pathlib import Path

REPO = Path(__file__).resolve().parent.parent
SRC_DIR = REPO / "docs" / "handoff-package"
OUT_DIR = SRC_DIR / "_bundle"

def _ordered_md_files() -> list[Path]:
    """Numeric prefix → ordered markdown sources. 00-* first, then 01-08."""
    return sorted(p for p in SRC_DIR.glob("*.md") if not p.name.startswith("_"))

# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def _build_docx(out_path: Path) -> None:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Cover
    cover = doc.add_paragraph()
    run = cover.add_run("[PLATFORM]\nArchitecture Handoff Package")
    run.bold = True
    run.font.size = Pt(28)
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(
        "Author: Trenton, Cloud Application Architect (Principal)\n"
        "Version 1.0 · 2026-05-07"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    fence_re = re.compile(r"^```(\w*)\s*$")
    table_sep_re = re.compile(r"^\s*\|?\s*[-:|\s]+\s*\|?\s*$")

    for md in _ordered_md_files():
        text = md.read_text(encoding="utf-8")
        in_fence = False
        fence_lang = ""
        fence_buf: list[str] = []
        table_buf: list[list[str]] = []

        def _flush_table() -> None:
            nonlocal table_buf
            if not table_buf:
                return
            cols = max(len(row) for row in table_buf)
            tbl = doc.add_table(rows=len(table_buf), cols=cols)
            tbl.style = "Light Grid Accent 1"
            for ri, row in enumerate(table_buf):
                for ci in range(cols):
                    cell = tbl.cell(ri, ci)
                    cell.text = row[ci] if ci < len(row) else ""
                    if ri == 0:
                        for r in cell.paragraphs[0].runs:
                            r.bold = True
            table_buf = []

        def _flush_code() -> None:
            nonlocal fence_buf, fence_lang
            if not fence_buf:
                return
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            run = p.add_run("\n".join(fence_buf))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            # Render mermaid/plantuml fences as a placeholder stamp + the source.
            if fence_lang in ("mermaid", "plantuml"):
                marker = doc.add_paragraph()
                m_run = marker.add_run(f"[Diagram: {fence_lang}]")
                m_run.italic = True
                m_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            fence_buf = []

        for raw in text.splitlines():
            line = raw.rstrip()

            fm = fence_re.match(line)
            if fm:
                if in_fence:
                    _flush_code()
                    in_fence = False
                else:
                    _flush_table()
                    in_fence = True
                    fence_lang = fm.group(1) or ""
                continue
            if in_fence:
                fence_buf.append(line)
                continue

            if line.startswith("|") and "|" in line[1:]:
                if not table_sep_re.match(line):
                    cells = [c.strip() for c in line.strip("|").split("|")]
                    table_buf.append(cells)
                continue
            else:
                _flush_table()

            if not line:
                doc.add_paragraph()
                continue

            if line.startswith("# "):
                doc.add_heading(line[2:].strip(), level=0)
            elif line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=1)
            elif line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=2)
            elif line.startswith("#### "):
                doc.add_heading(line[5:].strip(), level=3)
            elif line.startswith(("- ", "* ")):
                doc.add_paragraph(line[2:].strip(), style="List Bullet")
            elif re.match(r"^\d+\.\s", line):
                doc.add_paragraph(re.sub(r"^\d+\.\s", "", line), style="List Number")
            else:
                doc.add_paragraph(line)

        _flush_code()
        _flush_table()
        doc.add_page_break()

    doc.save(str(out_path))

# ---------------------------------------------------------------------------
# Per-document PDF — uses weasyprint over an HTML render of the markdown.
# Mermaid/PlantUML fences are kept as `<pre>` blocks (image rendering is
# done in chatmode itself per the capstone test); the PDF carries the
# diagram source, the chatmode artifact is the rendered visual.
# ---------------------------------------------------------------------------

def _markdown_to_html(md_text: str) -> str:
    try:
        from markdown_it import MarkdownIt  # type: ignore
        md = MarkdownIt("commonmark", {"html": False, "breaks": True}).enable("table")
        return md.render(md_text)
    except ImportError:
        # Fallback: minimal escape + <pre>
        from html import escape

        return f"<pre>{escape(md_text)}</pre>"

_HTML_SHELL = """<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8"/>
<title>{title}</title>
<style>
  @page {{ size: Letter; margin: 0.7in; }}
  body {{ font-family: 'Calibri', sans-serif; font-size: 11pt; color: #1a1a1a; }}
  h1, h2, h3, h4 {{ color: #0a2540; page-break-after: avoid; }}
  h1 {{ font-size: 22pt; border-bottom: 2px solid #0a2540; padding-bottom: 4px; }}
  h2 {{ font-size: 16pt; margin-top: 18pt; }}
  h3 {{ font-size: 13pt; margin-top: 12pt; }}
  pre {{ background: #f4f6f8; padding: 8px 12px; border-left: 3px solid #2563eb;
         font-family: 'Consolas', 'Menlo', monospace; font-size: 9pt;
         white-space: pre-wrap; word-break: break-word; }}
  code {{ background: #eef2f7; padding: 1px 4px; border-radius: 3px; font-size: 9.5pt; }}
  pre code {{ background: transparent; padding: 0; }}
  table {{ border-collapse: collapse; margin: 8pt 0; width: 100%; }}
  th, td {{ border: 1px solid #cbd5e1; padding: 6px 9px; text-align: left;
            font-size: 10pt; vertical-align: top; }}
  th {{ background: #e2e8f0; }}
  blockquote {{ border-left: 3px solid #94a3b8; margin: 8pt 0; padding-left: 10pt;
                color: #475569; }}
  .cover {{ text-align: center; padding: 25vh 0; }}
  .cover h1 {{ border: none; font-size: 36pt; }}
  .cover .sub {{ font-size: 14pt; color: #475569; }}
</style>
</head>
<body>
{body}
</body>
</html>
"""

def _build_pdfs(out_dir: Path) -> list[Path]:
    from weasyprint import HTML  # type: ignore

    out_dir.mkdir(parents=True, exist_ok=True)
    written: list[Path] = []
    for md in _ordered_md_files():
        body_html = _markdown_to_html(md.read_text(encoding="utf-8"))
        html = _HTML_SHELL.format(title=md.stem, body=body_html)
        pdf_path = out_dir / f"{md.stem}.pdf"
        HTML(string=html, base_url=str(SRC_DIR)).write_pdf(str(pdf_path))
        written.append(pdf_path)
    return written

# ---------------------------------------------------------------------------
# Index page
# ---------------------------------------------------------------------------

def _build_index(pdfs: list[Path], docx_path: Path, out_dir: Path) -> Path:
    rows = "\n".join(
        f'  <li><a href="{p.name}">{p.stem}.pdf</a></li>' for p in pdfs
    )
    html = f"""<!doctype html>
<html><head><meta charset="utf-8"/><title>[PLATFORM] Architecture Handoff</title>
<style>body{{font-family:Calibri,sans-serif;max-width:780px;margin:40px auto;padding:0 20px;}}
h1{{color:#0a2540;}}li{{margin:6px 0;}}.pkg{{background:#f4f6f8;padding:10px 14px;border-left:3px solid #2563eb;margin:18px 0;}}</style></head>
<body>
<h1>[PLATFORM] Architecture Handoff Package</h1>
<p>Author: Trenton, Cloud Application Architect (Principal) · Version 1.0 · 2026-05-07</p>
<div class="pkg">
  <strong>Bundle:</strong> {docx_path.name} (combined Word document)<br/>
  <a href="{docx_path.name}">Download {docx_path.name}</a>
</div>
<h2>Per-document PDFs</h2>
<ul>
{rows}
</ul>
</body></html>
"""
    idx = out_dir / "index.html"
    idx.write_text(html, encoding="utf-8")
    return idx

def main() -> int:
    if not SRC_DIR.exists():
        print(f"[bundle] source dir missing: {SRC_DIR}", file=sys.stderr)
        return 2
    md_files = _ordered_md_files()
    if not md_files:
        print(f"[bundle] no markdown files in {SRC_DIR}", file=sys.stderr)
        return 2

    OUT_DIR.mkdir(parents=True, exist_ok=True)

    print(f"[bundle] {len(md_files)} markdown files → {OUT_DIR}")
    print("[bundle] writing per-document PDFs (weasyprint)…")
    pdfs = _build_pdfs(OUT_DIR)
    for p in pdfs:
        print(f"  · {p.relative_to(REPO)}")

    docx_path = OUT_DIR / "platform-architecture-handoff-package.docx"
    print(f"[bundle] writing combined DOCX (python-docx) → {docx_path.relative_to(REPO)}")
    _build_docx(docx_path)

    idx = _build_index(pdfs, docx_path, OUT_DIR)
    print(f"[bundle] index → {idx.relative_to(REPO)}")
    print("[bundle] done.")
    return 0

if __name__ == "__main__":
    raise SystemExit(main())
